from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer


ROOT = Path(__file__).resolve().parent
DOCX_PATH = ROOT / "CS329E_Final_Report_Group23.docx"
PDF_PATH = ROOT / "CS329E_Final_Report_Group23.pdf"


TITLE = "Forecasting Airline Stock Volatility with HAR-RV, IV, OVX, and TOSI"
GROUP_NUMBER = "23"
NAMES_AND_IDS = (
    "Blake Stanley (bks2356), "
    "Shivsagar Palla (sp56633), "
    "Raghuvendra Chowdhry (rbc993)"
)

ANALYSIS_TEXT = (
    "We did not make any changes to our analysis after the presentation. The main "
    "analytical change happened after Phase 2: we moved away from the earlier "
    "MIDAS/monthly forecasting plan and rebuilt the project around a HAR-style "
    "realized-volatility framework. In the final version, we computed realized "
    "volatility from 5-minute Alpaca data, combined HAR-RV features with IV, OVX, "
    "and TOSI, and evaluated OLS, Ridge, Random Forest, and XGBoost with walk-forward "
    "out-of-sample testing. This gave us a cleaner time-series baseline, more direct "
    "comparisons across feature sets, and a stronger connection between forecasting "
    "results and the trading strategy."
)

DESIGN_BULLETS = [
    "Visualization packages used: Plotly for interactive charts and Streamlit for the dashboard interface.",
    "Interactive daily realized-volatility line charts with range sliders, shaded periods, and hover tooltips.",
    "An oil-driver view that overlays airline volatility with OVX and TOSI so users can inspect co-movement over time.",
    "Monthly correlation heatmaps and scatterplot facets for IV, OVX, and TOSI against future airline volatility.",
    "Model-comparison visuals for RMSE, Sharpe ratio, and Diebold-Mariano significance across feature sets and model families.",
    "Forecast-vs-actual panels, cumulative JETS straddle P&L charts, and feature-importance bars for the tree-based models.",
]

STRENGTHS_BULLETS = [
    "The project has a clear economic story: oil-market uncertainty and sentiment should matter for airline volatility.",
    "We used multiple data sources and frequencies, including intraday realized volatility, options-implied volatility, OVX, and TOSI.",
    "The HAR baseline made the modeling framework interpretable, and the feature-addition design let us test what each signal contributed.",
    "Our evaluation was disciplined: walk-forward out-of-sample testing, per-ticker comparisons, and trading-metric validation.",
    "The final Streamlit app made the project easier to explain because users could interact with the evidence instead of only seeing static charts.",
]

CHALLENGES_BULLETS = [
    "Aligning mixed-frequency data was difficult because realized volatility and IV are daily/intraday while TOSI is monthly.",
    "Several datasets had missing or uneven coverage, especially older implied-volatility history, which required careful cleaning and alignment.",
    "The external oil signals were informative, but improvements were not uniform across every airline, so interpretation had to stay nuanced.",
    "Running walk-forward experiments across several model families and feature sets was computationally expensive and required saving artifacts carefully.",
]

ADVICE_TEXT = (
    "Start the data pipeline early and lock down your target variable before you build models or visuals. "
    "Choose a strong, interpretable baseline first, because it becomes much easier to justify added complexity later. "
    "Also separate expensive modeling from the presentation layer by saving clean artifacts for the dashboard. "
    "Finally, do not rely only on correlations or in-sample fit; use true out-of-sample testing and keep your narrative tied to a concrete decision problem."
)


def set_run_font(run, bold=False):
    run.bold = bold
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(12)


def style_paragraph(paragraph, center=False):
    fmt = paragraph.paragraph_format
    fmt.line_spacing = 1.5
    fmt.space_after = Pt(0)
    fmt.space_before = Pt(0)
    if center:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_label_value(document, label, value):
    p = document.add_paragraph()
    style_paragraph(p)
    set_run_font(p.add_run(label), bold=True)
    set_run_font(p.add_run(value))


def add_section_paragraph(document, heading, text):
    p = document.add_paragraph()
    style_paragraph(p)
    set_run_font(p.add_run(heading), bold=True)
    set_run_font(p.add_run(text))


def add_bullets(document, bullets):
    for bullet in bullets:
        p = document.add_paragraph(style="List Bullet")
        style_paragraph(p)
        set_run_font(p.add_run(bullet))


def build_docx():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)

    title = doc.add_paragraph()
    style_paragraph(title, center=True)
    set_run_font(title.add_run(TITLE), bold=True)

    add_label_value(doc, "Group Number: ", GROUP_NUMBER)
    add_label_value(doc, "Your Names and UT IDS: ", NAMES_AND_IDS)
    add_section_paragraph(doc, "Analysis: ", ANALYSIS_TEXT)

    p = doc.add_paragraph()
    style_paragraph(p)
    set_run_font(p.add_run("Design:"), bold=True)
    add_bullets(doc, DESIGN_BULLETS)

    p = doc.add_paragraph()
    style_paragraph(p)
    set_run_font(p.add_run("Strengths:"), bold=True)
    add_bullets(doc, STRENGTHS_BULLETS)

    p = doc.add_paragraph()
    style_paragraph(p)
    set_run_font(p.add_run("Challenges:"), bold=True)
    add_bullets(doc, CHALLENGES_BULLETS)

    add_section_paragraph(doc, "Advice: ", ADVICE_TEXT)
    doc.save(DOCX_PATH)


def export_pdf():
    fonts_dir = Path(r"C:\Windows\Fonts")
    pdfmetrics.registerFont(TTFont("TimesNewRoman", str(fonts_dir / "times.ttf")))
    pdfmetrics.registerFont(TTFont("TimesNewRoman-Bold", str(fonts_dir / "timesbd.ttf")))

    doc = SimpleDocTemplate(
        str(PDF_PATH),
        pagesize=letter,
        leftMargin=inch,
        rightMargin=inch,
        topMargin=inch,
        bottomMargin=inch,
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "ReportTitle",
        parent=styles["Normal"],
        fontName="TimesNewRoman-Bold",
        fontSize=12,
        leading=18,
        alignment=TA_CENTER,
        spaceAfter=0,
    )
    body_style = ParagraphStyle(
        "ReportBody",
        parent=styles["Normal"],
        fontName="TimesNewRoman",
        fontSize=12,
        leading=18,
        spaceAfter=0,
    )
    bullet_style = ParagraphStyle(
        "ReportBullet",
        parent=body_style,
        leftIndent=18,
        firstLineIndent=0,
        bulletIndent=6,
    )

    story = [
        Paragraph(TITLE, title_style),
        Paragraph(f"<b>Group Number:</b> {GROUP_NUMBER}", body_style),
        Paragraph(f"<b>Your Names and UT IDS:</b> {NAMES_AND_IDS}", body_style),
        Paragraph(f"<b>Analysis:</b> {ANALYSIS_TEXT}", body_style),
        Paragraph("<b>Design:</b>", body_style),
    ]

    for bullet in DESIGN_BULLETS:
        story.append(Paragraph(bullet, bullet_style, bulletText="•"))

    story.append(Paragraph("<b>Strengths:</b>", body_style))
    for bullet in STRENGTHS_BULLETS:
        story.append(Paragraph(bullet, bullet_style, bulletText="•"))

    story.append(Paragraph("<b>Challenges:</b>", body_style))
    for bullet in CHALLENGES_BULLETS:
        story.append(Paragraph(bullet, bullet_style, bulletText="•"))

    story.extend(
        [
            Paragraph(f"<b>Advice:</b> {ADVICE_TEXT}", body_style),
            Spacer(1, 0),
        ]
    )
    doc.build(story)


if __name__ == "__main__":
    build_docx()
    export_pdf()
    print(f"Wrote {DOCX_PATH.name}")
    print(f"Wrote {PDF_PATH.name}")
